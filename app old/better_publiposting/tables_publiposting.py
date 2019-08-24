import docx
from typing import Dict, List
from .interfaces import Table, Document
from .docx_template import TableData
import copy

def render_table(table: Table, data:TableData):
    print(data.title)
    base_style = table.rows[0].cells[0].paragraphs[0].style.paragraph_format
    print(dir(base_style))
    for row in data.rows:
        _row = table.add_row()
        for cell, d in zip(_row.cells, row):
            # p = cell.add_paragraph()
            # p.style = base_style
            cell.paragraphs[0].style = None
            cell.paragraphs[0].text = d
            
            
    # first we add the rows with their content
    


def render_tables(doc: Document, data: Dict[str, TableData]) -> None:
    """Should be able to build tables from pre-existing ones
    """
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, _data in data.items():
                        if key in paragraph.text:
                            render_table(t, data[key])
