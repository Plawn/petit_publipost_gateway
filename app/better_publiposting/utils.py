import docx
from docx import Document
from typing import Dict, List, Generator, Iterable
import re
from lxml import etree
import json


# interfaces
class Run:
    def __init__(self):
        self.text: str = ''


class Paragraph:
    def __init__(self):
        self.runs: List[Run] = None
        self.text = ''


def find_all(a_str: str, sub: str) -> Generator[str, None, None]:
    start = 0
    while True:
        start = a_str.find(sub, start)
        if start == -1:
            return
        yield start
        start += len(sub)  # use start += 1 to find overlapping matches


def get_tables_paragraph(doc: Document) -> Generator[Paragraph, None, None]:
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def get_headers_paragraphs(doc: Document) -> Generator[Paragraph, None, None]:
    return []


def get_all_paragraphs(doc: Document) -> Generator[Paragraph, None, None]:
    # return [*get_tables_paragraph(doc), *list(doc.paragraphs), *get_headers_paragraphs(doc)]
    for p in get_tables_paragraph(doc):
        yield p
    for p in list(doc.paragraphs):
        yield p
    for p in get_headers_paragraphs(doc):
        yield p


def render(paragraph: Paragraph, key: str, value: str) -> None:
    _map = [0] * len(paragraph.text)
    offset = 0
    for m, run in enumerate(paragraph.runs):
        l = len(run.text)
        for k in range(l):
            _map[offset + k] = m
        offset += l

    for i in find_all(paragraph.text, key):
        remaining = len(key)
        l = list()
        while remaining > 0:
            k = len(paragraph.runs[_map[i]].text)
            remaining -= k
            l.append((_map[i], k))
            i += k

        # replacing start on key wirh the complete value
        i, _len = l[0]
        t = paragraph.runs[i].text
        paragraph.runs[i].text = t.replace(key[:_len], value)

        # making the key smaller as we already replaced a part of it
        # blanking the rest of the key
        key = key[_len:]
        for i, _len in l[1:]:
            t = paragraph.runs[i].text
            paragraph.runs[i].text = t.replace(key[:_len], '')
            key = key[_len:]


def docx_replace(doc: Document,  data: Dict[str, str]) -> None:
    """
    Makes text replacement in the whole document
    """
    paragraphs = get_all_paragraphs(doc)

    # das a placeholder for tests
    for paragraph in paragraphs:
        for key, value in data.items():
            key_name = f'{{{{{key}}}}}'
            if key_name in paragraph.text:
                runs = paragraph.runs
                for i in range(len(runs)):
                    if key_name in runs[i].text:
                        t: str = runs[i].text
                        runs[i].text = t.replace(key_name, value)
                        continue
                    render(paragraph, key_name, value)
                    break


def xml_cleaner(words: Iterable) -> Generator[str, None, None]:
    """Enlève les tags XML résiduels pour une liste de mots"""
    for word in words:
        cleaned_word = ""
        in_tag = False
        for j in range(len(word)):
            if word[j] == "<":
                in_tag = True
            elif word[j] == ">":
                in_tag = False
            elif not in_tag:
                cleaned_word += word[j]
        yield cleaned_word
